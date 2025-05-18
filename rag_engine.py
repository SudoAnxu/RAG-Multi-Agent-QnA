import os
import faiss
from sentence_transformers import SentenceTransformer
import pdfplumber
import docx

class RAGEngine:
    def __init__(self, docs_folder=None, embedding_model='all-MiniLM-L6-v2', chunk_size=350):
        self.embedding_model = embedding_model
        self.chunk_size = chunk_size
        self.encoder = SentenceTransformer(self.embedding_model)
        self.chunks = []
        self.sources = []
        self.index = None

        if docs_folder:
            self.ingest_local_files(docs_folder)

    def ingest_local_files(self, folder_path):
        # Read all supported files from the local folder and ingest them
        self.chunks, self.sources = [], []
        for filename in os.listdir(folder_path):
            if not filename.lower().endswith(('.txt', '.pdf', '.docx')):
                continue
            file_path = os.path.join(folder_path, filename)
            with open(file_path, "rb") as f:
                # Use the same ingestion code but from local file bytes
                if filename.endswith(".txt"):
                    text = f.read().decode("utf-8")
                elif filename.endswith(".pdf"):
                    text = self.extract_text_from_pdf(f)
                elif filename.endswith(".docx"):
                    text = self.extract_text_from_docx(f)
                else:
                    continue

                self._chunk_text(text, filename)
        self._build_faiss_index()

    def ingest_and_chunk_uploaded_files(self, uploaded_files):
        # For uploaded files from Streamlit uploader
        self.chunks, self.sources = [], []
        for file in uploaded_files:
            name = file.name
            if name.endswith(".txt"):
                text = file.read().decode("utf-8")
            elif name.endswith(".pdf"):
                text = self.extract_text_from_pdf(file)
            elif name.endswith(".docx"):
                text = self.extract_text_from_docx(file)
            else:
                continue
            self._chunk_text(text, name)
        self._build_faiss_index()

    def _chunk_text(self, text, source_name):
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            para = para.strip()
            for i in range(0, len(para), self.chunk_size):
                chunk = para[i:i+self.chunk_size].strip()
                if chunk:
                    self.chunks.append(chunk)
                    self.sources.append(source_name)

    def _build_faiss_index(self):
        if self.chunks:
            embeds = self.encoder.encode(self.chunks).astype("float32")
            dim = embeds.shape[1]
            self.index = faiss.IndexFlatL2(dim)
            self.index.add(embeds)

    def extract_text_from_pdf(self, file):
        import io
        if hasattr(file, "read"):
            # Reset file pointer
            file.seek(0)
            pdf = pdfplumber.open(file)
        else:
            pdf = pdfplumber.open(io.BytesIO(file))
        text = "\n".join([page.extract_text() or "" for page in pdf.pages])
        pdf.close()
        return text

    def extract_text_from_docx(self, file):
        import io
        if hasattr(file, "read"):
            file.seek(0)
            doc = docx.Document(file)
        else:
            doc = docx.Document(io.BytesIO(file))
        return "\n".join([para.text for para in doc.paragraphs])

    def retrieve(self, query, k=3):
        if not self.index:
            return []
        q_emb = self.encoder.encode([query]).astype("float32")
        D, I = self.index.search(q_emb, k)
        return [(self.chunks[i], self.sources[i]) for i in I[0]]
